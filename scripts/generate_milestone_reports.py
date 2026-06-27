"""Generate filled Milestone 1 and Milestone 2 Word reports.

Builds ``Milestone_1_Insurance_Price_Prediction.docx`` and
``Milestone_2_Insurance_Price_Prediction.docx`` at the repo root, following the
structure of the provided templates in ``templates/`` and embedding the figures
and tables produced by the pipeline (``reports/``) and the trained model
(``models/metrics.json``).

Run:  uv run python scripts/generate_milestone_reports.py
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "reports" / "figures"
TAB = ROOT / "reports" / "tables"
METRICS = json.loads((ROOT / "models" / "metrics.json").read_text())

INDIGO = RGBColor(0x4F, 0x46, 0xE5)


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def title_page(doc: Document, milestone: str) -> None:
    for text, size, bold, color in [
        (f"{milestone} Submission", 22, True, INDIGO),
        ("Insurance Price Prediction", 18, True, None),
        ("", 8, False, None),
        ("Submitted By", 12, True, None),
        ("Group No. << n >>   [Batch: << year >>]", 11, False, None),
        ("Group Members: << Name >>, << Name >>, << Name >>, << Name >>", 11, False, None),
        ("Mentor: << Name of mentor >>", 11, False, None),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
    doc.add_page_break()


def contents(doc: Document, sections: list[str]) -> None:
    h = doc.add_heading("Contents", level=1)
    for i, s in enumerate(sections, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number" if False else "Normal")
    doc.add_page_break()


def heading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def figure(doc: Document, name: str, caption: str, width: float = 6.0) -> None:
    path = FIG / name
    if not path.exists():
        para(doc, f"[figure {name} not found — run the notebooks to generate it]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure: {caption}")
    run.italic = True
    run.font.size = Pt(9)


def table_from_df(doc: Document, df: pd.DataFrame, max_rows: int | None = None) -> None:
    if max_rows:
        df = df.head(max_rows)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    for j, col in enumerate(df.columns):
        t.rows[0].cells[j].paragraphs[0].add_run(str(col)).bold = True
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            cells[j].text = f"{val:,.3f}" if isinstance(val, float) else str(val)


# --------------------------------------------------------------------------- #
# Milestone 1
# --------------------------------------------------------------------------- #
def build_milestone1() -> None:
    doc = Document()
    title_page(doc, "Milestone 1")
    contents(doc, [
        "Introduction", "Objectives, Problem Statement, and Sub-objectives",
        "Data Report", "Initial Exploratory Data Analysis", "Data Pre-processing",
        "Extensive Exploratory Data Analysis", "Appendix",
    ])

    heading(doc, "Introduction")
    para(doc,
         "Healthcare is one of the most important domains in the market because it is directly "
         "linked to an individual's life and finances. Treatment can be extremely costly, and an "
         "uninsured individual can face severe financial hardship. Medical insurers, in turn, want "
         "to optimise and price their policies to manage risk, because a healthy lifestyle "
         "substantially reduces the chance of illness. This project builds a data-driven model that "
         "estimates the optimum insurance cost for an individual using health, lifestyle, habit, and "
         "demographic parameters, supporting fair pricing, risk triage, and wellness interventions.")

    heading(doc, "Objectives, Problem Statement, and Sub-objectives")
    para(doc,
         "Problem statement: build a regression model that provides the optimum insurance cost for "
         "an individual from health and habit related parameters. The target variable is "
         "insurance_cost. Sub-objectives toward this goal:")
    bullets(doc, [
        "Load and audit the dataset; understand shape, types, and quality.",
        "Correct schema and category quality issues without losing business meaning.",
        "Treat missing values and outliers using defensible techniques.",
        "Engineer features capturing admission recency, BMI/cholesterol risk, disease history, and activity.",
        "Use EDA to identify the important predictors of insurance cost.",
        "Build a leak-free preprocessing pipeline (scaling/encoding fit on training data only).",
        "(Milestone 2) Build, tune, and select the best model, deploy it, and derive business insights.",
    ])

    heading(doc, "Data Report")
    para(doc,
         "The dataset (Insurance Data.csv) contains 25,000 rows and 24 columns — one row per "
         "applicant (applicant_id is unique). There are no duplicate rows. The target insurance_cost "
         "has a mean of ~27,147 and a median of ~27,148 with a mild right skew (~0.33). Only two "
         "columns contain missing values: Year_last_admitted (11,881 missing, 47.52%) and bmi "
         "(990 missing, 3.96%).")
    para(doc, "Variable types:")
    bullets(doc, [
        "Continuous: age, bmi, avg_glucose_level, daily_avg_steps, weight, fat_percentage, insurance_cost (target).",
        "Discrete/count: years_of_insurance_with_us, regular_checkup_last_year, visited_doctor_last_1_year, weight_change_in_last_one_year.",
        "Binary: adventure_sports, heart_disease_history, other_major_disease_history.",
        "Nominal categorical: Occupation, Gender, smoking_status, Location, covered_by_any_other_company, Alcohol, exercise.",
        "Ordinal categorical: cholesterol_level (banded ranges).",
        "Identifier (dropped): applicant_id.",
    ])
    para(doc,
         "Columns renamed to fix source typos: regular_checkup_lasy_year -> regular_checkup_last_year, "
         "heart_decs_history -> heart_disease_history, other_major_decs_history -> "
         "other_major_disease_history. The Occupation category 'Salried' was corrected to 'Salaried'.")
    if (TAB / "data_dictionary.csv").exists():
        para(doc, "Data dictionary (dtype, missing %, unique values):")
        dd = pd.read_csv(TAB / "data_dictionary.csv")
        dd.columns = ["column"] + list(dd.columns[1:])
        table_from_df(doc, dd)

    heading(doc, "Initial Exploratory Data Analysis")
    para(doc,
         "The response variable is insurance_cost; all remaining columns (except applicant_id) are "
         "candidate predictors. Univariate analysis of the continuous attributes shows roughly "
         "symmetric distributions for weight, bmi, glucose, and steps, while the target is mildly "
         "right-skewed. Categorical attributes are reasonably balanced: Occupation is dominated by "
         "Student/Business (~40% each), ~66% of applicants are Male, and ~30% have an 'Unknown' "
         "smoking status (retained as a valid business category).")
    figure(doc, "univariate_continuous.png", "Distributions and spread of continuous attributes.", 5.5)
    figure(doc, "univariate_categorical.png", "Category frequencies of categorical attributes.", 5.5)
    figure(doc, "missing_values.png", "Missing values by column (only Year_last_admitted and bmi).", 4.5)
    para(doc,
         "Missing-value decision: Year_last_admitted is missing for ~47.5% of applicants, which is "
         "meaningful — it indicates no prior hospital admission rather than a random gap. It is "
         "therefore re-expressed as engineered features (was_admitted_before, "
         "years_since_last_admitted) instead of being dropped or naively imputed. bmi has only ~4% "
         "missing and is imputed with the median by (age band, gender) with a global fallback, "
         "preserving its distribution.")

    heading(doc, "Data Pre-processing")
    bullets(doc, [
        "Dropped applicant_id (a unique identifier with no predictive signal); confirmed zero duplicate rows.",
        "Corrected the column-name typos and the 'Salried' category typo described above.",
        "bmi: imputed by (age_band, Gender) median with global-median fallback, then capped to a "
        "plausible 12-60 range to remove unrealistic extremes.",
        "Year_last_admitted: converted to was_admitted_before (0/1) and years_since_last_admitted "
        "(missing coded as no prior admission, recency 0).",
        "Outliers: reviewed via the IQR rule; target outliers were retained (high premiums can be "
        "valid and business-relevant) and only unrealistic bmi values were capped.",
        "Engineered features: age_band, bmi_category, cholesterol_midpoint, was_admitted_before, "
        "years_since_last_admitted, any_major_disease_history, weight_bmi_interaction, steps_per_age.",
        "Scaling: numeric features standardised inside a ColumnTransformer and categorical features "
        "one-hot encoded; all transformations are fit on the training split only to avoid leakage.",
    ])

    heading(doc, "Extensive Exploratory Data Analysis")
    para(doc,
         "Relationships with the target reveal that weight has an unusually strong positive "
         "association with insurance_cost, making it the single dominant driver in this dataset. "
         "(This is flagged for a plausibility/leakage review before any real-world deployment; it is "
         "retained in the model but documented.) Other notable signals: applicants covered by another "
         "insurer and those involved in adventure sports show clearly higher median costs, and prior "
         "admission history adds useful risk information.")
    figure(doc, "correlation_heatmap.png", "Correlation among numeric and engineered features.", 5.5)
    figure(doc, "scatter_vs_target.png", "Insurance cost versus key numeric drivers.", 6.0)
    figure(doc, "boxplots_by_category.png", "Insurance cost distribution across key categories.", 6.0)
    para(doc,
         "Summary: after schema fixes the data is clean, with only two columns needing missing-value "
         "treatment. EDA identifies weight, admission history, other-insurer coverage, and preventive "
         "checkups as the strongest cost signals. The preprocessing and feature-engineering pipeline "
         "produced here is reusable for the Milestone 2 modeling stage.")

    heading(doc, "Appendix")
    para(doc,
         "The complete, reproducible analysis is in experiments/01_milestone1_eda_preprocessing.ipynb, "
         "with all figures saved under reports/figures/ and the full data dictionary under "
         "reports/tables/data_dictionary.csv. Preprocessing and feature engineering are implemented as "
         "reusable scikit-learn transformers in src/insurance/ (data.py, features.py, preprocess.py).")

    out = ROOT / "Milestone_1_Insurance_Price_Prediction.docx"
    doc.save(out)
    print("Saved", out.name)


# --------------------------------------------------------------------------- #
# Milestone 2
# --------------------------------------------------------------------------- #
def build_milestone2() -> None:
    m = METRICS["metrics"]
    final = METRICS["final_model"]
    doc = Document()
    title_page(doc, "Milestone 2")
    contents(doc, [
        "Introduction", "Model Selection and Metrics of Interest",
        "Model Building and Evaluation", "Model Comparison and Selection",
        "Business Insights and Recommendations", "Appendix",
    ])

    heading(doc, "Introduction")
    para(doc,
         "This milestone builds on the cleaned data and engineered features from Milestone 1. The "
         "data is split into independent variables and the target insurance_cost using an 80/20 "
         "train-test split (random_state=42). A baseline model is established, a broad set of "
         "advanced models is trained and evaluated on both train and test sets, the strongest family "
         "is hyperparameter-tuned, the best model is selected, explained, and finally deployed via a "
         "FastAPI service, a React inference UI, and a Streamlit app.")

    heading(doc, "Model Selection and Metrics of Interest")
    para(doc, "Baseline models:")
    bullets(doc, [
        "DummyRegressor (predicts the mean) — establishes the no-skill floor.",
        "LinearRegression — first interpretable real baseline.",
    ])
    para(doc, "Advanced models:")
    bullets(doc, [
        "Regularized linear: Ridge, Lasso, ElasticNet.",
        "Trees and ensembles: DecisionTree, RandomForest, ExtraTrees.",
        "Gradient boosting: GradientBoosting, HistGradientBoosting, XGBoost, LightGBM, CatBoost.",
    ])
    para(doc, "Metrics of interest (regression):")
    bullets(doc, [
        "MAE — average absolute pricing error in cost units (primary, business-readable).",
        "RMSE — penalises larger pricing errors more heavily.",
        "R² — share of cost variance explained.",
        "MAPE / SMAPE — percentage-style errors for business communication.",
        "Cross-validated RMSE and the train-test RMSE gap — to check stability and overfitting.",
    ])

    heading(doc, "Model Building and Evaluation")
    para(doc,
         "Every model is wrapped in the same leak-free scikit-learn Pipeline (feature engineering -> "
         "preprocessing -> estimator); preprocessing is fit only on the training fold. Each model is "
         "evaluated on both the train and test sets, with 5-fold cross-validation for stability. "
         "Hyperparameter tuning uses RandomizedSearchCV (25 iterations, 5-fold CV, optimising "
         "negative RMSE) on the strongest family, with a finer Optuna study and a log-target "
         "transform experiment explored in the notebooks. Gradient boosting with XGBoost and CatBoost "
         "is GPU-accelerated; LightGBM runs multi-core on CPU.")
    if (TAB / "model_comparison.csv").exists():
        comp = pd.read_csv(TAB / "model_comparison.csv")
        comp = comp[["model", "test_MAE", "test_RMSE", "test_R2", "train_R2", "cv_rmse_mean", "train_test_RMSE_gap"]]
        para(doc, "Model evaluation (sorted by test RMSE):")
        table_from_df(doc, comp)
    figure(doc, "model_comparison.png", "Test RMSE by model (lower is better).", 5.5)

    heading(doc, "Model Comparison and Selection")
    para(doc,
         f"The gradient-boosting models clearly outperform the linear and single-tree models. "
         f"{final} was selected as the final model: it achieves the lowest test RMSE and MAE with a "
         f"strong R², a stable cross-validation score, and a small train-test gap (low "
         f"overfitting). RandomizedSearch tuning did not improve on the well-chosen defaults, so the "
         f"default configuration was retained for simplicity and robustness.")
    para(doc, "Final model performance on the held-out test set:")
    bullets(doc, [
        f"MAE: {m['test_MAE']:,.0f}",
        f"RMSE: {m['test_RMSE']:,.0f}",
        f"R²: {m['test_R2']:.4f}",
        f"MAPE: {m['test_MAPE']:.1f}%",
    ])
    figure(doc, "diagnostics.png", "Predicted vs actual and residuals for the final model.", 6.0)

    heading(doc, "Business Insights and Recommendations")
    para(doc,
         "Model explainability (permutation importance in production, corroborated by SHAP in the "
         "notebook) identifies the key drivers of predicted insurance cost: weight is by far the "
         "dominant factor, followed by admission history (Year_last_admitted), coverage by another "
         "insurer, preventive checkups, and weight change.")
    figure(doc, "feature_importance.png", "Permutation importance — top model drivers.", 5.0)
    figure(doc, "shap_summary.png", "SHAP summary — per-feature contribution and direction.", 5.0)
    figure(doc, "decile_gains.png", "Lift and cumulative gains by predicted-cost decile.", 6.0)
    para(doc, "Actionable recommendations (with measurable outcomes):")
    bullets(doc, [
        "Use the model as pricing decision-support, not automated underwriting — track quote "
        "turnaround time and underwriter override rate.",
        "Route extreme predictions (top decile) to manual review — measure reduction in mispriced "
        "high-cost policies.",
        "Target wellness programmes (weight management, preventive checkups, activity) at high-risk "
        "segments — track engagement and downstream claim reduction.",
        "Trigger documentation/preventive checks on prior-admission applicants — measure claim "
        "accuracy improvement.",
        "Apply fairness governance to Gender/Location and validate the dominance of weight for "
        "plausibility and leakage before production; monitor segment-level error and drift.",
    ])

    heading(doc, "Appendix")
    rt = METRICS.get("risk_thresholds", {})
    para(doc,
         "Deployment: the final scikit-learn Pipeline is saved to models/final_model.pkl and served "
         "by a FastAPI API (api/main.py, endpoints /health, /schema, /predict), a React + TypeScript "
         "inference UI (frontend/), and a Streamlit app (streamlit_app.py). Predictions are mapped to "
         "Low / Medium / High risk bands using training-set tertiles "
         f"(Low<{rt.get('low_to_medium', 0):,.0f}, High>{rt.get('medium_to_high', 0):,.0f}).")
    para(doc,
         "Reproducibility: run `uv run python -m insurance.train` to regenerate the model, metrics, "
         "and report tables. The full modeling workflow is in experiments/02-04, and the model "
         "comparison and importance tables are under reports/tables/.")
    para(doc,
         "Limitations: dataset provenance and real pricing rules are unknown; the dominance of weight "
         "must be validated for plausibility/leakage; all predictors must be available at quote time; "
         "fairness testing should be expanded before real-world deployment.")

    out = ROOT / "Milestone_2_Insurance_Price_Prediction.docx"
    doc.save(out)
    print("Saved", out.name)


if __name__ == "__main__":
    build_milestone1()
    build_milestone2()
